#!/usr/bin/env python3
"""Generate Playwright CI Analysis Report as DOCX"""

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

_REPO_ROOT = Path(__file__).resolve().parent.parent
_DEFAULT_OUTPUT = _REPO_ROOT / "docs" / "reports" / "playwright-ci-analysis-pr1214.docx"

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════

title = doc.add_heading('Playwright CI Analysis Report — PR #1214', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('SAHOOL National Agricultural Intelligence Platform')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
subtitle.add_run('\n')
run2 = subtitle.add_run('تقرير تحليل اختبارات Playwright في بيئة CI')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# BASIC INFO TABLE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Basic Information | المعلومات الأساسية', level=1)

info_data = [
    ('Project | المشروع', 'kafaat/sahool-unified-v15-idp'),
    ('PR', '#1214 — test(mobile): setup test database with Drift generated classes'),
    ('Commit', 'e1f88e5 (Merge commit) — 2026-03-13'),
    ('Playwright Result | نتيجة Playwright', 'PASSED (1/1) — misleading result (see C-1)'),
    ('Files Changed | الملفات المعدّلة', '26 files'),
    ('Date | التاريخ', '2026-03-14'),
]

table = doc.add_table(rows=len(info_data), cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (key, val) in enumerate(info_data):
    row = table.rows[i]
    row.cells[0].text = key
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = val

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# CRITICAL ISSUES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Critical Issues | مشاكل حرجة', level=1)

# C-1
h = doc.add_heading('C-1: Empty Playwright Test Suite | مجموعة اختبارات Playwright فارغة', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('CRITICAL')
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

doc.add_paragraph(
    'The only registered test file is no-tests-to-run.spec.ts containing a single skipped test '
    'with the message "skipped - backend API not available in CI". Duration: 3ms (Before Hooks + After Hooks only). '
    'No real E2E test is executed, making the CI check valueless — it always passes regardless of code changes.'
)

doc.add_paragraph(
    'الملف الوحيد المسجَّل هو no-tests-to-run.spec.ts ويحتوي على اختبار متجاهَل (skip) بعنوان '
    '"skipped - backend API not available in CI". المدة: 3 ميلي ثانية. لا يوجد أي اختبار E2E حقيقي، '
    'مما يجعل الـ CI check بلا قيمة.'
)

p = doc.add_paragraph()
p.add_run('Recommended Fix: ').bold = True
p.add_run(
    'Add Playwright tests that work without a backend (using MSW or mock server), '
    'or rename the CI check to "playwright-stub" to clearly indicate no real tests run.'
)

doc.add_paragraph()

# C-2
h = doc.add_heading('C-2: Certificate Pinning Removed | إزالة Certificate Pinning', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
p.add_run('Severity: ').bold = True
run = p.add_run('CRITICAL (Security)')
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

doc.add_paragraph(
    '87 lines defining certificate pinning for production domains were deleted from '
    'network_security_config.xml. The justification is valid (placeholder pins were causing '
    'connection rejection), but the result is: TLS without pinning = vulnerable to MITM attacks.'
)

doc.add_paragraph(
    'تم حذف 87 سطراً كانت تُعرّف certificate pinning لنطاقات الإنتاج من network_security_config.xml. '
    'السبب مُبرَّر (pins كانت placeholder تُسبب رفض الاتصالات)، لكن النتيجة: TLS بدون pinning = عرضة لهجمات MITM.'
)

p = doc.add_paragraph()
p.add_run('Recommended Fix: ').bold = True
p.add_run(
    'Replace placeholder pins with real SHA-256 pins from production certificates. '
    'Add pin rotation strategy with backup pins.'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# HIGH SEVERITY ISSUES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('High Severity Issues | مشاكل عالية الأهمية', level=1)

# H-1
h = doc.add_heading('H-1: Incorrect Order of libsqlite3-dev Installation in test.yml', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

doc.add_paragraph(
    'The libsqlite3-dev installation step was added after widget tests and before integration tests. '
    'If any widget test uses Drift/SQLite, it will fail because the native library is not yet available.'
)

doc.add_paragraph(
    'الخطوة أُضيفت بعد widget tests وقبل integration tests. إذا كانت أي widget test تستخدم Drift/SQLite ستفشل.'
)

p = doc.add_paragraph()
p.add_run('Recommended Fix: ').bold = True
p.add_run('Move the libsqlite3-dev installation step before all test steps in the CI workflow.')

doc.add_paragraph()

# H-2
h = doc.add_heading('H-2: pubspec.lock Changed Without Dependency Documentation', level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)

doc.add_paragraph(
    'pubspec.lock changed with +33/-9 lines without documenting which new dependencies were added. '
    'This makes it harder to review supply chain changes and potential security implications.'
)

p = doc.add_paragraph()
p.add_run('Recommended Fix: ').bold = True
p.add_run('Add a summary of dependency changes in the PR description.')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# MEDIUM SEVERITY ISSUES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Medium Severity Issues | مشاكل متوسطة', level=1)

medium_issues = [
    (
        'M-1: _consecutiveFailures reset timing',
        '_consecutiveFailures = 0 is reset after backoff timer expires, not after actual successful sync. '
        'This means the counter resets even if the next sync also fails.',
        '_consecutiveFailures = 0 يُصفَّر بعد انتهاء وقت الـ backoff، لا عند النجاح الفعلي.'
    ),
    (
        'M-2: SecureStorage._read swallows errors silently',
        'SecureStorage._read catches exceptions and returns null without notifying monitoring systems. '
        'This can hide persistent storage corruption issues.',
        'SecureStorage._read تبتلع الخطأ صامتةً وتُعيد null بدلاً من إشعار monitoring.'
    ),
    (
        'M-3: WeChat service deprecated without migration guide',
        'The WeChat service (wechat-service) was deprecated on port 8133 but no migration guide '
        'was provided for existing consumers.',
        'خدمة WeChat أُعلن إهمالها دون توفير دليل ترحيل.'
    ),
]

for title_text, desc_en, desc_ar in medium_issues:
    h = doc.add_heading(title_text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0xFF, 0xCC, 0x00)
    doc.add_paragraph(desc_en)
    p = doc.add_paragraph()
    run = p.add_run(desc_ar)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# POSITIVE CHANGES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Excellent Improvements in This PR | تحسينات ممتازة', level=1)

improvements = [
    ('Security: Certificate bypass limited to localhost in debug mode',
     'إصلاح أمني ممتاز: Certificate bypass محدود لـ localhost فقط في debug mode'),
    ('Resilience: Backoff + Jitter in sync_engine.dart',
     'Backoff + Jitter لتجنب thundering herd'),
    ('Memory Leak Fix: _networkSubscription?.cancel() before re-subscribing',
     'إصلاح تسريب ذاكرة: إلغاء الاشتراك القديم قبل إعادة الاشتراك'),
    ('Lifecycle: _disposed flag in auth_service.dart replaces mounted check',
     '_disposed flag بدلاً من mounted'),
    ('DRY: _serviceUrl(devPort) replaces 143 repeated lines in env_config.dart',
     'تجفيف الكود: _serviceUrl(devPort) يستبدل 143 سطر متكرر'),
    ('Testing: Comprehensive Drift tests with in-memory database + composite indexes',
     'اختبارات Drift شاملة مع قاعدة بيانات في الذاكرة'),
    ('Feature: Community service (Rocket.Chat) integrated on port 8133 with 8 endpoints',
     'خدمة Community مُدمجة على port 8133 مع 8 endpoints'),
]

for en, ar in improvements:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(en)
    run.bold = True
    p.add_run('\n')
    run2 = p.add_run(ar)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run2.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Issues Summary | ملخص المشاكل', level=1)

summary_data = [
    ('Severity', 'Count', 'IDs'),
    ('Critical', '2', 'C-1, C-2'),
    ('High', '2', 'H-1, H-2'),
    ('Medium', '3', 'M-1, M-2, M-3'),
    ('Improvements', '7', '—'),
]

table = doc.add_table(rows=len(summary_data), cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (sev, count, ids) in enumerate(summary_data):
    row = table.rows[i]
    for j, val in enumerate([sev, count, ids]):
        row.cells[j].text = val
        if i == 0:
            row.cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading('Recommendations | التوصيات', level=1)

recommendations = [
    'Add real Playwright E2E tests or mock-based tests that validate UI rendering without backend dependency.',
    'Replace placeholder certificate pins with production SHA-256 pins before next mobile release.',
    'Move libsqlite3-dev installation before widget tests in CI workflow.',
    'Document dependency changes in PR descriptions for supply chain visibility.',
    'Fix _consecutiveFailures reset to only occur on successful sync.',
    'Add monitoring/alerting for SecureStorage read failures.',
    'Create migration guide for WeChat service deprecation.',
]

for i, rec in enumerate(recommendations, 1):
    doc.add_paragraph(f'{i}. {rec}')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Generated: 2026-03-14 | SAHOOL Platform v16.0.0 | KAFAAT')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
parser = argparse.ArgumentParser(description="Generate Playwright CI Analysis Report")
parser.add_argument(
    "-o", "--output",
    type=Path,
    default=_DEFAULT_OUTPUT,
    help=f"Output file path (default: {_DEFAULT_OUTPUT.relative_to(_REPO_ROOT)})",
)
args = parser.parse_args()

output_path = Path(args.output)
output_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(output_path))
print(f'Report saved to: {output_path}')
